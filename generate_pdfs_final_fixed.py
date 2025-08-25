#!/usr/bin/env python3
"""
Fixed PDF Generation Script with proper stamp paper matching and formatting preservation
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
import glob
import subprocess
import time
import difflib

def find_matching_stamp_paper(name, stamp_papers):
    """Find the stamp paper that best matches the shareholder name"""
    # Clean the name for matching
    clean_name = str(name).strip().upper()
    
    # Try exact match first (case-insensitive)
    for stamp_path in stamp_papers:
        stamp_filename = os.path.basename(stamp_path)
        stamp_name = os.path.splitext(stamp_filename)[0].upper()
        
        if clean_name == stamp_name:
            return stamp_path
    
    # Try partial match - if stamp paper name contains shareholder name
    for stamp_path in stamp_papers:
        stamp_filename = os.path.basename(stamp_path)
        stamp_name = os.path.splitext(stamp_filename)[0].upper()
        
        if clean_name in stamp_name or stamp_name in clean_name:
            return stamp_path
    
    # Try fuzzy matching to find closest match
    stamp_names = []
    for stamp_path in stamp_papers:
        stamp_filename = os.path.basename(stamp_path)
        stamp_name = os.path.splitext(stamp_filename)[0]
        stamp_names.append((stamp_name, stamp_path))
    
    # Find closest match using difflib
    matches = difflib.get_close_matches(name, [sn[0] for sn in stamp_names], n=1, cutoff=0.6)
    if matches:
        for stamp_name, stamp_path in stamp_names:
            if stamp_name == matches[0]:
                return stamp_path
    
    # If no match found, return None
    print(f"    ⚠ No matching stamp paper found for: {name}")
    return None

def replace_text_in_runs(paragraph, replacements):
    """Replace text while preserving formatting in runs"""
    for key, value in replacements.items():
        if key in paragraph.text:
            # Work with runs to preserve formatting
            inline = paragraph.runs
            for i in range(len(inline)):
                if key in inline[i].text:
                    # Preserve the run's formatting
                    text = inline[i].text.replace(key, str(value))
                    inline[i].text = text

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables while preserving formatting"""
    try:
        doc = Document(doc_path)
        
        # Replace in paragraphs while preserving formatting
        for paragraph in doc.paragraphs:
            if paragraph.text:
                replace_text_in_runs(paragraph, replacements)
        
        # Replace in tables while preserving formatting
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            replace_text_in_runs(paragraph, replacements)
        
        # Replace in headers
        for section in doc.sections:
            header = section.header
            if header:
                for paragraph in header.paragraphs:
                    if paragraph.text:
                        replace_text_in_runs(paragraph, replacements)
        
        # Replace in footers
        for section in doc.sections:
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    if paragraph.text:
                        replace_text_in_runs(paragraph, replacements)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"    Error: {e}")
        return False

def convert_docx_to_pdf_libreoffice(docx_path, output_dir):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        cmd = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            return True
        else:
            print(f"    LibreOffice error: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        print("    Conversion timed out")
        return False
    except Exception as e:
        print(f"    Conversion error: {e}")
        return False

def process_record(row, documents_dir, stamp_papers, output_folder, temp_dir, index, total):
    """Process a single shareholder record"""
    
    # Get shareholder details
    name = row.get('Name of the Shareholder', 'Unknown')
    entity_type = row.get('Entity', 'Individual')
    email = row.get('Email', '')
    address = row.get('Address', '')
    
    # Clean entity type
    if pd.isna(entity_type) or entity_type == '':
        entity_type = 'Individual'
    
    print(f"\n[{index}/{total}] Processing: {name}")
    print(f"    Type: {entity_type}")
    
    # Select template
    if entity_type.lower() == 'entity':
        template_path = os.path.join(documents_dir, 'Entity.docx')
    else:
        template_path = os.path.join(documents_dir, 'Individual.docx')
    
    if not os.path.exists(template_path):
        print(f"    ✗ Template not found: {template_path}")
        return False
    
    # Create safe filename
    safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()
    if len(safe_name) > 50:
        safe_name = safe_name[:50]
    
    # Prepare replacements
    current_day = datetime.now().day
    replacements = {
        '[Name of the Shareholder]': str(name),
        '[Email]': str(email),
        '[Address]': str(address),
        '[Date]': str(current_day),
        '[date]': str(current_day)  # Handle lowercase variant
    }
    
    # Create temporary files
    temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
    temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
    
    # Step 1: Replace variables
    print("    Replacing variables (preserving formatting)...")
    if not replace_variables_in_doc(template_path, replacements, temp_docx):
        print("    ✗ Failed to replace variables")
        return False
    
    # Step 2: Convert to PDF
    print("    Converting to PDF...")
    if not convert_docx_to_pdf_libreoffice(temp_docx, temp_dir):
        print("    ✗ Failed to convert to PDF")
        return False
    
    # Wait for PDF to be created
    time.sleep(0.5)
    
    if not os.path.exists(temp_pdf):
        print("    ✗ PDF file not created")
        return False
    
    # Step 3: Find matching stamp paper
    matching_stamp = find_matching_stamp_paper(name, stamp_papers)
    
    # Step 4: Merge with stamp paper
    final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
    
    if matching_stamp:
        print(f"    Merging with stamp paper: {os.path.basename(matching_stamp)}")
        try:
            merger = PdfMerger()
            
            # Add matching stamp paper first
            merger.append(matching_stamp)
            
            # Add the document
            merger.append(temp_pdf)
            
            # Write final PDF
            merger.write(final_pdf)
            merger.close()
            
            print(f"    ✓ Success: {os.path.basename(final_pdf)}")
            return True
            
        except Exception as e:
            print(f"    ✗ Merge failed: {e}")
            return False
    else:
        # If no stamp paper found, just save the document PDF
        print("    ⚠ Creating PDF without stamp paper")
        shutil.copy(temp_pdf, final_pdf)
        print(f"    ✓ Success (no stamp): {os.path.basename(final_pdf)}")
        return True

def main():
    """Main function"""
    
    print("=" * 70)
    print(" " * 10 + "PDF GENERATION WITH MATCHING STAMP PAPERS")
    print("=" * 70)
    
    # Setup paths
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs_Matched")
    
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Check required files
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found at {excel_path}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Individual.docx')):
        print(f"Error: Individual.docx not found in {documents_dir}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Entity.docx')):
        print(f"Error: Entity.docx not found in {documents_dir}")
        return
    
    # Read data
    print("\nReading Excel file...")
    df = pd.read_excel(excel_path)
    total_records = len(df)
    
    # Get stamp papers
    stamp_papers = []
    if os.path.exists(stamp_folder):
        stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
        print(f"✓ Found {len(stamp_papers)} stamp papers")
        
        # Show sample stamp paper names
        print("\nSample stamp paper files:")
        for sp in stamp_papers[:5]:
            print(f"  - {os.path.basename(sp)}")
    else:
        print("⚠ No stamp papers folder found")
    
    print(f"✓ Found {total_records} records in Excel")
    
    # Get processing range
    print("\n" + "-" * 70)
    start = int(input(f"Start from record number (1-{total_records}): ") or "1") - 1
    count = int(input(f"How many records to process (default: 10): ") or "10")
    end = min(start + count, total_records)
    
    print(f"\nWill process records {start+1} to {end}")
    confirm = input("Continue? (y/n): ")
    if confirm.lower() != 'y':
        print("Cancelled")
        return
    
    print("\n" + "=" * 70)
    print("PROCESSING WITH MATCHED STAMP PAPERS...")
    print("=" * 70)
    
    # Create temporary directory
    temp_dir = tempfile.mkdtemp()
    success_count = 0
    no_stamp_count = 0
    failed_names = []
    
    try:
        for i in range(start, end):
            row = df.iloc[i]
            result = process_record(
                row,
                documents_dir,
                stamp_papers,
                output_folder,
                temp_dir,
                i - start + 1,
                end - start
            )
            
            if result:
                success_count += 1
            else:
                name = row.get('Name of the Shareholder', f'Record_{i+1}')
                failed_names.append(name)
            
            # Small delay between conversions
            time.sleep(0.5)
    
    except KeyboardInterrupt:
        print("\n\nInterrupted by user")
    
    finally:
        # Cleanup
        print("\nCleaning up temporary files...")
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    # Summary
    print("\n" + "=" * 70)
    print(" " * 25 + "SUMMARY")
    print("=" * 70)
    print(f"Total processed: {end - start}")
    print(f"✓ Successful: {success_count}")
    print(f"✗ Failed: {end - start - success_count}")
    
    if failed_names:
        print(f"\nFailed records:")
        for name in failed_names[:5]:
            print(f"  - {name}")
        if len(failed_names) > 5:
            print(f"  ... and {len(failed_names) - 5} more")
    
    print(f"\nOutput folder: {output_folder}")
    print("Note: Files are saved in 'Generated_PDFs_Matched' folder")
    
    if success_count > 0:
        print(f"\n✓ Successfully generated {success_count} PDFs with matched stamp papers!")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)