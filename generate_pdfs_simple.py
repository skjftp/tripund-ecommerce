#!/usr/bin/env python3
"""
Simplified PDF generation script - processes in batches
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
from pathlib import Path
import glob
import random
import subprocess

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document"""
    try:
        doc = Document(doc_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, str(value))
                            inline[i].text = text
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, str(value))
                                        inline[i].text = text
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error: {e}")
        return False

def convert_docx_to_pdf_libreoffice(docx_path, output_dir):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        cmd = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return True
        else:
            print(f"LibreOffice error: {result.stderr}")
            return False
    except Exception as e:
        print(f"Conversion error: {e}")
        return False

def process_batch(start_idx, end_idx):
    """Process a batch of shareholders"""
    
    # Configuration
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs")
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Read data
    df = pd.read_excel(excel_path)
    batch_df = df.iloc[start_idx:end_idx]
    
    # Get stamp papers
    stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
    
    # Create temp directory
    temp_dir = tempfile.mkdtemp()
    
    success_count = 0
    
    try:
        for idx, row in batch_df.iterrows():
            # Get details
            name = row.get('Name of the Shareholder', 'Unknown')
            entity_type = row.get('Entity', 'Individual')
            email = row.get('Email', '')
            address = row.get('Address', '')
            
            if pd.isna(entity_type) or entity_type == '':
                entity_type = 'Individual'
            
            # Select template
            if entity_type.lower() == 'entity':
                template_path = os.path.join(documents_dir, 'Entity.docx')
            else:
                template_path = os.path.join(documents_dir, 'Individual.docx')
            
            # Create safe filename
            safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()[:50]
            
            # Prepare replacements
            replacements = {
                '[Name of the Shareholder]': str(name),
                '[Email]': str(email),
                '[Address]': str(address),
                '[Date]': str(datetime.now().day)
            }
            
            # Create temp docx
            temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
            
            # Replace variables
            if not replace_variables_in_doc(template_path, replacements, temp_docx):
                print(f"✗ Failed to process {name}")
                continue
            
            # Convert to PDF
            if not convert_docx_to_pdf_libreoffice(temp_docx, temp_dir):
                print(f"✗ Failed to convert {name}")
                continue
            
            # Get the converted PDF
            temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
            
            if not os.path.exists(temp_pdf):
                print(f"✗ PDF not created for {name}")
                continue
            
            # Merge with stamp paper
            final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
            
            try:
                merger = PdfMerger()
                
                # Add random stamp paper
                if stamp_papers:
                    stamp_paper = random.choice(stamp_papers)
                    merger.append(stamp_paper)
                
                # Add the document
                merger.append(temp_pdf)
                
                # Save final PDF
                merger.write(final_pdf)
                merger.close()
                
                print(f"✓ Generated: {safe_name}_{entity_type}.pdf")
                success_count += 1
                
            except Exception as e:
                print(f"✗ Merge failed for {name}: {e}")
    
    finally:
        # Cleanup
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    return success_count

def main():
    """Main function"""
    
    print("=" * 60)
    print("PDF GENERATION SCRIPT")
    print("=" * 60)
    
    # Check LibreOffice
    libreoffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    if not os.path.exists(libreoffice_path):
        print("\nError: LibreOffice not found!")
        print("Please install LibreOffice from: https://www.libreoffice.org/download/")
        print("Or use: brew install --cask libreoffice")
        return
    
    # Process in batches
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    
    df = pd.read_excel(excel_path)
    total = len(df)
    
    print(f"\nProcessing {total} shareholders...")
    
    batch_size = int(input("\nEnter batch size (e.g., 10): ") or "10")
    start = int(input("Start from record number (1-based, default 1): ") or "1") - 1
    
    end = min(start + batch_size, total)
    
    print(f"\nProcessing records {start+1} to {end}...")
    
    success = process_batch(start, end)
    
    print(f"\n{'=' * 60}")
    print(f"Processed: {end - start} records")
    print(f"Successful: {success}")
    print(f"Failed: {end - start - success}")
    print(f"Output folder: ~/Documents/Projects/Generated_PDFs/")

if __name__ == "__main__":
    main()