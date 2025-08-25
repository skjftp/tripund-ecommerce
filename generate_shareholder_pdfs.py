#!/usr/bin/env python3
"""
Shareholder PDF Generation Script
Generates PDFs with stamp papers and personalized documents from Excel data
"""

import os
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from docx2pdf import convert
import tempfile
import shutil
from pathlib import Path
import glob
import random

def read_excel_data(excel_path):
    """Read and process the Excel file with shareholder information"""
    try:
        df = pd.read_excel(excel_path)
        print(f"Successfully loaded {len(df)} records from Excel")
        
        # Ensure required columns exist
        required_columns = ['First Name', 'Name of the Shareholder', 'Email', 'Entity', 'Address']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            print(f"Warning: Missing columns: {missing_columns}")
            print(f"Available columns: {df.columns.tolist()}")
        
        return df
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None

def get_stamp_paper(stamp_papers_folder):
    """Get a random stamp paper from the folder"""
    stamp_papers = glob.glob(os.path.join(stamp_papers_folder, "*.pdf"))
    if not stamp_papers:
        print(f"Warning: No PDF stamp papers found in {stamp_papers_folder}")
        return None
    return random.choice(stamp_papers)

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document"""
    try:
        doc = Document(doc_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Handle runs to preserve formatting
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))
        
        # Save the modified document
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error replacing variables in document: {e}")
        return False

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using python-docx2pdf or LibreOffice"""
    try:
        # Try using docx2pdf (Windows/Mac with MS Word)
        convert(docx_path, pdf_path)
        return True
    except:
        try:
            # Fallback to LibreOffice (Linux/Mac without MS Word)
            import subprocess
            cmd = [
                'soffice',
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                os.path.dirname(pdf_path),
                docx_path
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            
            # LibreOffice creates PDF with same name as input
            temp_pdf = os.path.join(
                os.path.dirname(pdf_path),
                os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            )
            if temp_pdf != pdf_path:
                shutil.move(temp_pdf, pdf_path)
            return True
        except Exception as e:
            print(f"Error converting DOCX to PDF: {e}")
            return False

def merge_pdfs(stamp_paper_path, document_pdf_path, output_path):
    """Merge stamp paper with the document PDF"""
    try:
        merger = PdfMerger()
        
        # Add stamp paper first (if exists)
        if stamp_paper_path and os.path.exists(stamp_paper_path):
            merger.append(stamp_paper_path)
        
        # Add the document
        if os.path.exists(document_pdf_path):
            merger.append(document_pdf_path)
        
        # Write the merged PDF
        merger.write(output_path)
        merger.close()
        return True
    except Exception as e:
        print(f"Error merging PDFs: {e}")
        return False

def process_shareholder(row, template_folder, stamp_papers_folder, output_folder, temp_dir):
    """Process a single shareholder record"""
    try:
        # Determine template type
        entity_type = row.get('Entity', 'Individual')
        if pd.isna(entity_type) or entity_type == '':
            entity_type = 'Individual'
        
        # Select appropriate template (handle case variations)
        if entity_type.lower() == 'entity':
            # Try both lowercase and capitalized versions
            template_file = os.path.join(template_folder, 'entity.docx')
            if not os.path.exists(template_file):
                template_file = os.path.join(template_folder, 'Entity.docx')
        else:
            # Try both lowercase and capitalized versions
            template_file = os.path.join(template_folder, 'individual.docx')
            if not os.path.exists(template_file):
                template_file = os.path.join(template_folder, 'Individual.docx')
        
        if not os.path.exists(template_file):
            print(f"Template not found: {template_file}")
            return False
        
        # Prepare replacements
        current_date = datetime.now().day  # Just the day number as requested
        
        # Get shareholder name - try different column names
        shareholder_name = row.get('Name of the Shareholder', '')
        if pd.isna(shareholder_name) or shareholder_name == '':
            shareholder_name = row.get('First Name', 'Unknown')
        
        replacements = {
            '[Name of the Shareholder]': shareholder_name,
            '[Email]': row.get('Email', ''),
            '[Address]': row.get('Address', ''),
            '[Date]': str(current_date)
        }
        
        # Clean filename
        safe_name = "".join(c for c in str(shareholder_name) if c.isalnum() or c in (' ', '-', '_')).rstrip()
        if not safe_name:
            safe_name = 'Unknown'
        
        # Create temporary files
        temp_docx = os.path.join(temp_dir, f"{safe_name}_temp.docx")
        temp_pdf = os.path.join(temp_dir, f"{safe_name}_temp.pdf")
        
        # Replace variables in document
        if not replace_variables_in_doc(template_file, replacements, temp_docx):
            return False
        
        # Convert to PDF
        if not convert_docx_to_pdf(temp_docx, temp_pdf):
            print(f"Failed to convert {safe_name} document to PDF")
            return False
        
        # Get stamp paper
        stamp_paper_path = get_stamp_paper(stamp_papers_folder)
        
        # Create final output path
        output_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
        
        # Merge PDFs
        if not merge_pdfs(stamp_paper_path, temp_pdf, output_pdf):
            return False
        
        print(f"✓ Generated PDF for {shareholder_name}: {output_pdf}")
        return True
        
    except Exception as e:
        print(f"Error processing shareholder: {e}")
        return False

def main():
    """Main function to orchestrate the PDF generation"""
    
    # Configuration - Files are in Documents/Projects folder
    documents_projects_dir = os.path.expanduser("~/Documents/Projects")
    
    # Input paths
    excel_path = os.path.join(documents_projects_dir, "Email + Address List.xlsx")
    template_folder = documents_projects_dir  # Templates are in Documents/Projects
    stamp_papers_folder = os.path.join(documents_projects_dir, "DOA Stamp Papers")
    
    # Output path (in the same Documents/Projects folder)
    output_folder = os.path.join(documents_projects_dir, "Generated_PDFs")
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    # Check if required files exist
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found at {excel_path}")
        print("Please ensure 'Email + Address List.xlsx' is in ~/Documents/Projects/")
        return
    
    # Check for templates (case-insensitive)
    individual_template = os.path.join(template_folder, 'individual.docx')
    if not os.path.exists(individual_template):
        individual_template = os.path.join(template_folder, 'Individual.docx')
    
    entity_template = os.path.join(template_folder, 'entity.docx')
    if not os.path.exists(entity_template):
        entity_template = os.path.join(template_folder, 'Entity.docx')
    
    if not os.path.exists(individual_template):
        print(f"Error: individual.docx or Individual.docx template not found in {template_folder}")
        print("Please add the Individual.docx template file to ~/Documents/Projects/")
        return
    
    if not os.path.exists(entity_template):
        print(f"Error: entity.docx or Entity.docx template not found in {template_folder}")
        print("Please add the Entity.docx template file to ~/Documents/Projects/")
        return
    
    if not os.path.exists(stamp_papers_folder):
        print(f"Warning: Stamp papers folder not found at {stamp_papers_folder}")
        print("PDFs will be generated without stamp papers")
        os.makedirs(stamp_papers_folder, exist_ok=True)
    
    # Read Excel data
    df = read_excel_data(excel_path)
    if df is None:
        return
    
    # Create temporary directory
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Process each shareholder
        success_count = 0
        total_count = len(df)
        
        print(f"\nProcessing {total_count} shareholders...")
        print("-" * 50)
        
        for index, row in df.iterrows():
            if process_shareholder(row, template_folder, stamp_papers_folder, output_folder, temp_dir):
                success_count += 1
        
        print("-" * 50)
        print(f"\nSummary:")
        print(f"  Total records: {total_count}")
        print(f"  Successfully generated: {success_count}")
        print(f"  Failed: {total_count - success_count}")
        print(f"\nPDFs saved in: {output_folder}")
        
    finally:
        # Clean up temporary directory
        shutil.rmtree(temp_dir, ignore_errors=True)

if __name__ == "__main__":
    print("=" * 60)
    print("SHAREHOLDER PDF GENERATION SCRIPT")
    print("=" * 60)
    print()
    main()