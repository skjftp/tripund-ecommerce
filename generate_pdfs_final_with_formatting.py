#!/usr/bin/env python3
"""
Final PDF Generation Script with proper formatting preservation
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
import glob
import subprocess
import time
import difflib
import re

def find_matching_stamp_paper(name, stamp_papers):
    """Find the stamp paper that best matches the shareholder name with fuzzy matching"""
    # Clean the name for matching
    clean_name = str(name).strip().upper()
    
    # Remove common suffixes for better matching
    suffixes = ['PRIVATE LIMITED', 'PVT LTD', 'LIMITED', 'LTD', 'LLP', 'HUF']
    clean_name_short = clean_name
    for suffix in suffixes:
        clean_name_short = clean_name_short.replace(suffix, '').strip()
    
    best_match = None
    best_score = 0
    
    for stamp_path in stamp_papers:
        stamp_filename = os.path.basename(stamp_path)
        stamp_name = os.path.splitext(stamp_filename)[0].upper()
        stamp_name_short = stamp_name
        
        # Remove suffixes from stamp name too
        for suffix in suffixes:
            stamp_name_short = stamp_name_short.replace(suffix, '').strip()
        
        # Calculate similarity score
        score = difflib.SequenceMatcher(None, clean_name_short, stamp_name_short).ratio()
        
        # Check for exact match first
        if clean_name == stamp_name:
            return stamp_path
        
        # Check if one contains the other
        if clean_name_short in stamp_name_short or stamp_name_short in clean_name_short:
            if len(clean_name_short) > 5 and len(stamp_name_short) > 5:  # Avoid short matches
                score = max(score, 0.9)
        
        # Keep track of best match
        if score > best_score:
            best_score = score
            best_match = stamp_path
    
    # Return best match if score is good enough
    if best_score > 0.8:  # 80% similarity threshold
        return best_match
    elif best_score > 0.7:
        print(f"    ⚠ Partial match (score: {best_score:.2f}): {os.path.basename(best_match)}")
        return best_match
    else:
        print(f"    ⚠ No good match found (best score: {best_score:.2f})")
        return None

def replace_text_preserve_char_format(paragraph, replacements):
    """Replace text in paragraph while preserving character-level formatting"""
    
    # Get the full text
    full_text = paragraph.text
    
    # Check if any replacements are needed
    needs_replacement = False
    for key in replacements.keys():
        if key in full_text:
            needs_replacement = True
            break
    
    if not needs_replacement:
        return
    
    # Build a map of character positions to run formatting
    char_formats = []
    current_pos = 0
    
    for run in paragraph.runs:
        run_text = run.text
        run_length = len(run_text)
        
        # Store formatting for each character in this run
        for i in range(run_length):
            char_formats.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
            })
        current_pos += run_length
    
    # Perform replacements in the text
    new_text = full_text
    offset = 0  # Track position changes due to replacements
    
    for key, value in replacements.items():
        # Find all occurrences of the key
        start = 0
        while True:
            pos = new_text.find(key, start)
            if pos == -1:
                break
            
            # Calculate the formatting adjustment
            len_diff = len(str(value)) - len(key)
            
            # Replace the text
            new_text = new_text[:pos] + str(value) + new_text[pos + len(key):]
            
            # Adjust formatting map
            # Remove formatting for the old key length
            for _ in range(len(key)):
                if pos < len(char_formats):
                    char_formats.pop(pos)
            
            # Insert formatting for the new value
            # Use the formatting from the first character of the replaced text
            if pos < len(char_formats):
                format_to_copy = char_formats[pos-1] if pos > 0 else char_formats[0] if char_formats else {'bold': None, 'italic': None, 'underline': None, 'font_name': None, 'font_size': None}
            else:
                format_to_copy = {'bold': None, 'italic': None, 'underline': None, 'font_name': None, 'font_size': None}
            
            for _ in range(len(str(value))):
                char_formats.insert(pos, format_to_copy.copy())
            
            start = pos + len(str(value))
    
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""
    
    # Rebuild runs with preserved formatting
    if not char_formats:
        # No formatting info, just set the text
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
    else:
        # Group consecutive characters with same formatting into runs
        current_run = None
        current_format = None
        current_text = ""
        
        for i, char in enumerate(new_text):
            if i < len(char_formats):
                char_format = char_formats[i]
            else:
                char_format = {'bold': None, 'italic': None, 'underline': None, 'font_name': None, 'font_size': None}
            
            # Check if format changed
            if char_format != current_format:
                # Save previous run if exists
                if current_text:
                    run = paragraph.add_run(current_text)
                    if current_format:
                        if current_format.get('bold'):
                            run.bold = True
                        if current_format.get('italic'):
                            run.italic = True
                        if current_format.get('underline'):
                            run.underline = True
                        if current_format.get('font_name'):
                            run.font.name = current_format.get('font_name')
                        if current_format.get('font_size'):
                            run.font.size = current_format.get('font_size')
                
                # Start new run
                current_format = char_format
                current_text = char
            else:
                current_text += char
        
        # Add final run
        if current_text:
            run = paragraph.add_run(current_text)
            if current_format:
                if current_format.get('bold'):
                    run.bold = True
                if current_format.get('italic'):
                    run.italic = True
                if current_format.get('underline'):
                    run.underline = True
                if current_format.get('font_name'):
                    run.font.name = current_format.get('font_name')
                if current_format.get('font_size'):
                    run.font.size = current_format.get('font_size')

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document while preserving mixed formatting"""
    try:
        doc = Document(doc_path)
        
        # Replace in paragraphs with format preservation
        for paragraph in doc.paragraphs:
            if paragraph.text:
                # Check if this paragraph has mixed formatting
                has_mixed_format = False
                if len(paragraph.runs) > 1:
                    # Check if different runs have different formatting
                    first_bold = paragraph.runs[0].bold
                    for run in paragraph.runs[1:]:
                        if run.bold != first_bold:
                            has_mixed_format = True
                            break
                
                if has_mixed_format:
                    # Use character-level preservation for mixed format paragraphs
                    replace_text_preserve_char_format(paragraph, replacements)
                else:
                    # Simple replacement for uniform format paragraphs
                    text = paragraph.text
                    changed = False
                    for key, value in replacements.items():
                        if key in text:
                            text = text.replace(key, str(value))
                            changed = True
                    
                    if changed:
                        if paragraph.runs:
                            # Preserve format from first run
                            first_run = paragraph.runs[0]
                            is_bold = first_run.bold
                            is_italic = first_run.italic
                            is_underline = first_run.underline
                            font_name = first_run.font.name
                            font_size = first_run.font.size
                            
                            # Clear runs and set new text
                            for run in paragraph.runs:
                                run.text = ""
                            first_run.text = text
                            
                            # Restore formatting
                            if is_bold:
                                first_run.bold = True
                            if is_italic:
                                first_run.italic = True
                            if is_underline:
                                first_run.underline = True
                            if font_name:
                                first_run.font.name = font_name
                            if font_size:
                                first_run.font.size = font_size
                        else:
                            paragraph.text = text
        
        # Replace in tables (simpler approach as tables usually have uniform formatting)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            text = paragraph.text
                            changed = False
                            for key, value in replacements.items():
                                if key in text:
                                    text = text.replace(key, str(value))
                                    changed = True
                            
                            if changed:
                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    is_bold = first_run.bold
                                    for run in paragraph.runs:
                                        run.text = ""
                                    first_run.text = text
                                    if is_bold:
                                        first_run.bold = True
                                else:
                                    paragraph.text = text
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"    Error: {e}")
        return False

def convert_docx_to_pdf_libreoffice(docx_path, output_dir):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        cmd = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            return True
        else:
            print(f"    LibreOffice error: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        print("    Conversion timed out")
        return False
    except Exception as e:
        print(f"    Conversion error: {e}")
        return False

def process_record(row, documents_dir, stamp_papers, output_folder, temp_dir, index, total):
    """Process a single shareholder record"""
    
    # Get shareholder details
    name = row.get('Name of the Shareholder', 'Unknown')
    entity_type = row.get('Entity', 'Individual')
    email = row.get('Email', '')
    address = row.get('Address', '')
    
    # Clean entity type
    if pd.isna(entity_type) or entity_type == '':
        entity_type = 'Individual'
    
    print(f"\n[{index}/{total}] Processing: {name}")
    print(f"    Type: {entity_type}")
    
    # Select template
    if entity_type.lower() == 'entity':
        template_path = os.path.join(documents_dir, 'Entity.docx')
    else:
        template_path = os.path.join(documents_dir, 'Individual.docx')
    
    if not os.path.exists(template_path):
        print(f"    ✗ Template not found: {template_path}")
        return False
    
    # Create safe filename
    safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()
    if len(safe_name) > 50:
        safe_name = safe_name[:50]
    
    # Prepare replacements
    current_day = datetime.now().day
    replacements = {
        '[Name of the Shareholder]': str(name),
        '[Email]': str(email),
        '[Address]': str(address),
        '[Date]': str(current_day),
        '[date]': str(current_day)
    }
    
    # Create temporary files
    temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
    temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
    
    # Step 1: Replace variables
    print("    Replacing variables (preserving formatting)...")
    if not replace_variables_in_doc(template_path, replacements, temp_docx):
        print("    ✗ Failed to replace variables")
        return False
    
    # Step 2: Convert to PDF
    print("    Converting to PDF...")
    if not convert_docx_to_pdf_libreoffice(temp_docx, temp_dir):
        print("    ✗ Failed to convert to PDF")
        return False
    
    # Wait for PDF to be created
    time.sleep(0.5)
    
    if not os.path.exists(temp_pdf):
        print("    ✗ PDF file not created")
        return False
    
    # Step 3: Find matching stamp paper
    matching_stamp = find_matching_stamp_paper(name, stamp_papers)
    
    # Step 4: Merge with stamp paper
    final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
    
    if matching_stamp:
        print(f"    Merging with: {os.path.basename(matching_stamp)}")
        try:
            merger = PdfMerger()
            merger.append(matching_stamp)
            merger.append(temp_pdf)
            merger.write(final_pdf)
            merger.close()
            
            print(f"    ✓ Success: {os.path.basename(final_pdf)}")
            return True
            
        except Exception as e:
            print(f"    ✗ Merge failed: {e}")
            return False
    else:
        print("    Creating PDF without stamp paper")
        shutil.copy(temp_pdf, final_pdf)
        print(f"    ✓ Success (no stamp): {os.path.basename(final_pdf)}")
        return True

def main():
    """Main function"""
    
    print("=" * 70)
    print(" " * 10 + "PDF GENERATION WITH PRESERVED FORMATTING")
    print("=" * 70)
    
    # Setup paths
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs_Final")
    
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Check required files
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found at {excel_path}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Individual.docx')):
        print(f"Error: Individual.docx not found in {documents_dir}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Entity.docx')):
        print(f"Error: Entity.docx not found in {documents_dir}")
        return
    
    # Read data
    print("\nReading Excel file...")
    df = pd.read_excel(excel_path)
    total_records = len(df)
    
    # Get stamp papers
    stamp_papers = []
    if os.path.exists(stamp_folder):
        stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
        print(f"✓ Found {len(stamp_papers)} stamp papers")
    
    print(f"✓ Found {total_records} records in Excel")
    
    # Get processing range
    print("\n" + "-" * 70)
    start = int(input(f"Start from record number (1-{total_records}): ") or "1") - 1
    count = int(input(f"How many records to process (default: ALL): ") or str(total_records - start))
    end = min(start + count, total_records)
    
    print(f"\nWill process records {start+1} to {end}")
    confirm = input("Continue? (y/n): ")
    if confirm.lower() != 'y':
        print("Cancelled")
        return
    
    print("\n" + "=" * 70)
    print("PROCESSING...")
    print("=" * 70)
    
    # Create temporary directory
    temp_dir = tempfile.mkdtemp()
    success_count = 0
    failed_names = []
    
    try:
        for i in range(start, end):
            row = df.iloc[i]
            result = process_record(
                row,
                documents_dir,
                stamp_papers,
                output_folder,
                temp_dir,
                i - start + 1,
                end - start
            )
            
            if result:
                success_count += 1
            else:
                name = row.get('Name of the Shareholder', f'Record_{i+1}')
                failed_names.append(name)
            
            time.sleep(0.5)
    
    except KeyboardInterrupt:
        print("\n\nInterrupted by user")
    
    finally:
        print("\nCleaning up temporary files...")
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    # Summary
    print("\n" + "=" * 70)
    print(" " * 25 + "SUMMARY")
    print("=" * 70)
    print(f"Total processed: {end - start}")
    print(f"✓ Successful: {success_count}")
    print(f"✗ Failed: {end - start - success_count}")
    
    if failed_names:
        print(f"\nFailed records:")
        for name in failed_names[:5]:
            print(f"  - {name}")
    
    print(f"\nOutput folder: {output_folder}")
    print("Note: Files are saved in 'Generated_PDFs_Final' folder")
    
    if success_count > 0:
        print(f"\n✓ Successfully generated {success_count} PDFs with preserved formatting!")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)