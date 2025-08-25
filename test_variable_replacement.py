#!/usr/bin/env python3
"""
Test script to debug variable replacement
"""

import os
import pandas as pd
from datetime import datetime
from docx import Document
import tempfile

# Setup paths
documents_dir = os.path.expanduser("~/Documents/Projects")
excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")

# Read first record
df = pd.read_excel(excel_path)
row = df.iloc[0]

# Get details
name = row.get('Name of the Shareholder', 'Unknown')
entity_type = row.get('Entity', 'Individual')
email = row.get('Email', '')
address = row.get('Address', '')

print("Data from Excel:")
print(f"  Name: {name}")
print(f"  Entity: {entity_type}")
print(f"  Email: {email}")
print(f"  Address: {address}")

# Select template
if entity_type.lower() == 'entity':
    template_path = os.path.join(documents_dir, 'Entity.docx')
else:
    template_path = os.path.join(documents_dir, 'Individual.docx')

print(f"\nUsing template: {template_path}")

# Open and check template content
doc = Document(template_path)

print("\nChecking template for variables...")
print("=" * 50)

# Check paragraphs
for i, paragraph in enumerate(doc.paragraphs):
    if '[' in paragraph.text and ']' in paragraph.text:
        print(f"Paragraph {i}: {paragraph.text[:100]}...")

# Check tables
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if '[' in paragraph.text and ']' in paragraph.text:
                    print(f"Table {t_idx}, Row {r_idx}, Cell {c_idx}: {paragraph.text[:100]}...")

print("\n" + "=" * 50)
print("Testing replacement...")

# Create temp file
temp_dir = tempfile.mkdtemp()
temp_docx = os.path.join(temp_dir, "test_output.docx")

# Prepare replacements
replacements = {
    '[Name of the Shareholder]': str(name),
    '[Email]': str(email),
    '[Address]': str(address),
    '[Date]': str(datetime.now().day)
}

print("\nReplacements to make:")
for key, value in replacements.items():
    print(f"  {key} -> {value}")

# Create new document for testing
doc = Document(template_path)

replaced_count = 0

# Replace in paragraphs
print("\nReplacing in paragraphs...")
for paragraph in doc.paragraphs:
    original_text = paragraph.text
    for key, value in replacements.items():
        if key in paragraph.text:
            print(f"  Found '{key}' in paragraph")
            # Try replacing in the full paragraph text
            paragraph.text = paragraph.text.replace(key, str(value))
            replaced_count += 1

# Replace in tables
print("\nReplacing in tables...")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        print(f"  Found '{key}' in table cell")
                        paragraph.text = paragraph.text.replace(key, str(value))
                        replaced_count += 1

# Save and check
doc.save(temp_docx)
print(f"\nTotal replacements made: {replaced_count}")
print(f"Test document saved to: {temp_docx}")

# Read it back to verify
doc2 = Document(temp_docx)
print("\nVerifying saved document...")
found_variables = False
for paragraph in doc2.paragraphs:
    if '[' in paragraph.text and ']' in paragraph.text:
        print(f"  Still has variables: {paragraph.text[:100]}...")
        found_variables = True

if not found_variables:
    print("  ✓ All variables appear to be replaced!")
else:
    print("  ✗ Some variables were not replaced")

# Cleanup
import shutil
shutil.rmtree(temp_dir)