#!/usr/bin/env python3
"""
Debug script to test variable replacement
"""

import os
from docx import Document
from datetime import datetime
import tempfile

# Setup
documents_dir = os.path.expanduser("~/Documents/Projects")
template_path = os.path.join(documents_dir, 'Entity.docx')

# Test data
name = "AAMARA CAPITAL PRIVATE LIMITED"
email = "test@example.com"
address = "Test Address"
current_day = datetime.now().day

print("Testing variable replacement...")
print("=" * 50)

# Open document
doc = Document(template_path)

# Check what's in the document
print("\nSearching for variables in document:")
found_vars = []
for paragraph in doc.paragraphs:
    if '[' in paragraph.text and ']' in paragraph.text:
        # Extract variables
        text = paragraph.text
        start = 0
        while '[' in text[start:]:
            idx_start = text.index('[', start)
            idx_end = text.index(']', idx_start) if ']' in text[idx_start:] else len(text)
            var = text[idx_start:idx_end+1]
            if var not in found_vars:
                found_vars.append(var)
                print(f"  Found: {var}")
            start = idx_end

print(f"\nTotal variables found: {len(found_vars)}")

# Now test replacement methods
print("\n" + "=" * 50)
print("Method 1: Direct paragraph.text replacement")
print("=" * 50)

doc1 = Document(template_path)
replacements = {
    '[Name of the Shareholder]': name,
    '[Email]': email,
    '[Address]': address,
    '[Date]': str(current_day),
    '[date]': str(current_day)
}

for paragraph in doc1.paragraphs:
    if paragraph.text:
        original = paragraph.text
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                print(f"  Replaced {key} in paragraph")

# Save and check
temp_file1 = tempfile.mktemp(suffix='.docx')
doc1.save(temp_file1)

# Check if replacements worked
doc1_check = Document(temp_file1)
remaining1 = 0
for paragraph in doc1_check.paragraphs:
    for key in replacements.keys():
        if key in paragraph.text:
            remaining1 += 1
            print(f"  WARNING: {key} still present!")

print(f"Result: {remaining1} variables remaining")

print("\n" + "=" * 50)
print("Method 2: Working with runs")
print("=" * 50)

doc2 = Document(template_path)

for paragraph in doc2.paragraphs:
    if paragraph.text:
        for key, value in replacements.items():
            if key in paragraph.text:
                # Check runs
                print(f"  Found {key} in paragraph with {len(paragraph.runs)} runs")
                for i, run in enumerate(paragraph.runs):
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        print(f"    Replaced in run {i}")
                    elif '[' in run.text or ']' in run.text:
                        print(f"    Run {i} has partial: '{run.text}'")

# Save and check
temp_file2 = tempfile.mktemp(suffix='.docx')
doc2.save(temp_file2)

doc2_check = Document(temp_file2)
remaining2 = 0
for paragraph in doc2_check.paragraphs:
    for key in replacements.keys():
        if key in paragraph.text:
            remaining2 += 1
            print(f"  WARNING: {key} still present!")

print(f"Result: {remaining2} variables remaining")

print("\n" + "=" * 50)
print("Method 3: Reconstruct paragraph from scratch")
print("=" * 50)

doc3 = Document(template_path)

for paragraph in doc3.paragraphs:
    if paragraph.text:
        text = paragraph.text
        changed = False
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, value)
                changed = True
                print(f"  Replacing {key}")
        
        if changed:
            # Keep first run's formatting
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                # Set new text in first run
                first_run.text = text
            else:
                paragraph.text = text

# Save and check
temp_file3 = tempfile.mktemp(suffix='.docx')
doc3.save(temp_file3)

doc3_check = Document(temp_file3)
remaining3 = 0
for paragraph in doc3_check.paragraphs:
    for key in replacements.keys():
        if key in paragraph.text:
            remaining3 += 1
            print(f"  WARNING: {key} still present!")

print(f"Result: {remaining3} variables remaining")

print("\n" + "=" * 50)
print("SUMMARY:")
print(f"  Method 1 (paragraph.text): {remaining1} vars remaining")
print(f"  Method 2 (runs): {remaining2} vars remaining")
print(f"  Method 3 (reconstruct): {remaining3} vars remaining")

# Cleanup
import os
os.unlink(temp_file1)
os.unlink(temp_file2)
os.unlink(temp_file3)