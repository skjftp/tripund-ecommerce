#!/usr/bin/env python3
"""
Final PDF Generation Script using LibreOffice
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
import glob
import random
import subprocess
import time

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document"""
    try:
        doc = Document(doc_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # Replace at paragraph level first
                        inline = paragraph.runs
                        # Start from scratch with the paragraph text
                        started = False
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                inline[i].text = inline[i].text.replace(key, str(value))
                                started = True
                            elif started:
                                # Sometimes Word splits the variable across runs
                                # Check if this run completes a partial match
                                pass
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        if key in inline[i].text:
                                            inline[i].text = inline[i].text.replace(key, str(value))
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"    Error: {e}")
        return False

def convert_docx_to_pdf_libreoffice(docx_path, output_dir):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        # Use the installed LibreOffice
        cmd = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            return True
        else:
            print(f"    LibreOffice error: {result.stderr}")
            return False
    except subprocess.TimeoutExpired:
        print("    Conversion timed out")
        return False
    except Exception as e:
        print(f"    Conversion error: {e}")
        return False

def process_record(row, documents_dir, stamp_papers, output_folder, temp_dir, index, total):
    """Process a single shareholder record"""
    
    # Get shareholder details
    name = row.get('Name of the Shareholder', 'Unknown')
    entity_type = row.get('Entity', 'Individual')
    email = row.get('Email', '')
    address = row.get('Address', '')
    
    # Clean entity type
    if pd.isna(entity_type) or entity_type == '':
        entity_type = 'Individual'
    
    print(f"\n[{index}/{total}] Processing: {name}")
    print(f"    Type: {entity_type}")
    
    # Select template
    if entity_type.lower() == 'entity':
        template_path = os.path.join(documents_dir, 'Entity.docx')
    else:
        template_path = os.path.join(documents_dir, 'Individual.docx')
    
    if not os.path.exists(template_path):
        print(f"    ✗ Template not found: {template_path}")
        return False
    
    # Create safe filename
    safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()
    if len(safe_name) > 50:
        safe_name = safe_name[:50]
    
    # Prepare replacements
    current_day = datetime.now().day
    replacements = {
        '[Name of the Shareholder]': str(name),
        '[Email]': str(email),
        '[Address]': str(address),
        '[Date]': str(current_day),
        '[date]': str(current_day)  # Handle both cases
    }
    
    # Create temporary files
    temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
    temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
    
    # Step 1: Replace variables
    print("    Replacing variables...")
    if not replace_variables_in_doc(template_path, replacements, temp_docx):
        print("    ✗ Failed to replace variables")
        return False
    
    # Step 2: Convert to PDF
    print("    Converting to PDF...")
    if not convert_docx_to_pdf_libreoffice(temp_docx, temp_dir):
        print("    ✗ Failed to convert to PDF")
        return False
    
    # Wait for PDF to be created
    time.sleep(0.5)
    
    if not os.path.exists(temp_pdf):
        print("    ✗ PDF file not created")
        return False
    
    # Step 3: Merge with stamp paper
    final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
    
    print("    Merging with stamp paper...")
    try:
        merger = PdfMerger()
        
        # Add random stamp paper first
        if stamp_papers:
            stamp_paper = random.choice(stamp_papers)
            merger.append(stamp_paper)
            print(f"    Using stamp: {os.path.basename(stamp_paper)[:30]}...")
        
        # Add the document
        merger.append(temp_pdf)
        
        # Write final PDF
        merger.write(final_pdf)
        merger.close()
        
        print(f"    ✓ Success: {os.path.basename(final_pdf)}")
        return True
        
    except Exception as e:
        print(f"    ✗ Merge failed: {e}")
        return False

def main():
    """Main function"""
    
    print("=" * 70)
    print(" " * 20 + "PDF GENERATION SCRIPT")
    print("=" * 70)
    
    # Setup paths
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs")
    
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Check required files
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found at {excel_path}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Individual.docx')):
        print(f"Error: Individual.docx not found in {documents_dir}")
        return
    
    if not os.path.exists(os.path.join(documents_dir, 'Entity.docx')):
        print(f"Error: Entity.docx not found in {documents_dir}")
        return
    
    # Read data
    print("\nReading Excel file...")
    df = pd.read_excel(excel_path)
    total_records = len(df)
    
    # Get stamp papers
    stamp_papers = []
    if os.path.exists(stamp_folder):
        stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
    
    print(f"✓ Found {total_records} records")
    print(f"✓ Found {len(stamp_papers)} stamp papers")
    
    # Get processing range
    print("\n" + "-" * 70)
    start = int(input(f"Start from record number (1-{total_records}): ") or "1") - 1
    count = int(input(f"How many records to process (default: 10): ") or "10")
    end = min(start + count, total_records)
    
    print(f"\nWill process records {start+1} to {end}")
    confirm = input("Continue? (y/n): ")
    if confirm.lower() != 'y':
        print("Cancelled")
        return
    
    print("\n" + "=" * 70)
    print("PROCESSING...")
    print("=" * 70)
    
    # Create temporary directory
    temp_dir = tempfile.mkdtemp()
    success_count = 0
    
    try:
        for i in range(start, end):
            if process_record(
                df.iloc[i],
                documents_dir,
                stamp_papers,
                output_folder,
                temp_dir,
                i - start + 1,
                end - start
            ):
                success_count += 1
            
            # Small delay between conversions
            time.sleep(0.5)
    
    except KeyboardInterrupt:
        print("\n\nInterrupted by user")
    
    finally:
        # Cleanup
        print("\nCleaning up temporary files...")
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    # Summary
    print("\n" + "=" * 70)
    print(" " * 25 + "SUMMARY")
    print("=" * 70)
    print(f"Total processed: {end - start}")
    print(f"✓ Successful: {success_count}")
    print(f"✗ Failed: {end - start - success_count}")
    print(f"\nOutput folder: {output_folder}")
    
    if success_count > 0:
        print(f"\n✓ Successfully generated {success_count} PDFs!")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)