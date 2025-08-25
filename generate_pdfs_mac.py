#!/usr/bin/env python3
"""
PDF Generation Script for Mac using docx2pdf
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
import glob
import random
from docx2pdf import convert

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document"""
    try:
        doc = Document(doc_path)
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
        
        # Replace in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error replacing variables: {e}")
        return False

def main():
    """Main function to process PDFs"""
    
    print("=" * 60)
    print("PDF GENERATION SCRIPT FOR MAC")
    print("=" * 60)
    
    # Setup paths
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs")
    
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Read Excel data
    print("\nReading Excel file...")
    df = pd.read_excel(excel_path)
    total = len(df)
    print(f"Found {total} records")
    
    # Get stamp papers
    stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
    print(f"Found {len(stamp_papers)} stamp papers")
    
    # Ask for processing range
    start = int(input(f"\nStart from record (1-{total}): ") or "1") - 1
    count = int(input(f"How many to process (max {total-start}): ") or "5")
    end = min(start + count, total)
    
    print(f"\nProcessing records {start+1} to {end}...")
    print("-" * 60)
    
    # Create temp directory
    temp_dir = tempfile.mkdtemp()
    success_count = 0
    
    try:
        for idx in range(start, end):
            row = df.iloc[idx]
            
            # Get shareholder details
            name = row.get('Name of the Shareholder', 'Unknown')
            entity_type = row.get('Entity', 'Individual')
            email = row.get('Email', '')
            address = row.get('Address', '')
            
            # Clean entity type
            if pd.isna(entity_type) or entity_type == '':
                entity_type = 'Individual'
            
            print(f"\n[{idx-start+1}/{count}] Processing: {name} ({entity_type})")
            
            # Select template
            if entity_type.lower() == 'entity':
                template_path = os.path.join(documents_dir, 'Entity.docx')
            else:
                template_path = os.path.join(documents_dir, 'Individual.docx')
            
            # Create safe filename
            safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()
            if len(safe_name) > 50:
                safe_name = safe_name[:50]
            
            # Prepare replacements
            current_day = datetime.now().day
            replacements = {
                '[Name of the Shareholder]': str(name),
                '[Email]': str(email),
                '[Address]': str(address),
                '[Date]': str(current_day)
            }
            
            # Create temporary DOCX
            temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
            temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
            
            # Replace variables
            print("  - Replacing variables...")
            if not replace_variables_in_doc(template_path, replacements, temp_docx):
                print("  ✗ Failed to replace variables")
                continue
            
            # Convert to PDF
            print("  - Converting to PDF...")
            try:
                convert(temp_docx, temp_pdf)
            except Exception as e:
                print(f"  ✗ PDF conversion failed: {e}")
                continue
            
            if not os.path.exists(temp_pdf):
                print("  ✗ PDF file not created")
                continue
            
            # Create final PDF with stamp paper
            final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
            
            print("  - Merging with stamp paper...")
            try:
                merger = PdfMerger()
                
                # Add random stamp paper first
                if stamp_papers:
                    stamp_paper = random.choice(stamp_papers)
                    merger.append(stamp_paper)
                    print(f"  - Using stamp: {os.path.basename(stamp_paper)}")
                
                # Add the document
                merger.append(temp_pdf)
                
                # Write final PDF
                merger.write(final_pdf)
                merger.close()
                
                print(f"  ✓ Success: {os.path.basename(final_pdf)}")
                success_count += 1
                
            except Exception as e:
                print(f"  ✗ Merge failed: {e}")
    
    except KeyboardInterrupt:
        print("\n\nInterrupted by user")
    
    finally:
        # Cleanup temp directory
        print("\nCleaning up temporary files...")
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total processed: {end - start}")
    print(f"Successful: {success_count}")
    print(f"Failed: {end - start - success_count}")
    print(f"\nOutput folder: {output_folder}")
    print(f"Files generated: {success_count} PDFs")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)