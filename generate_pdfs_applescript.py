#!/usr/bin/env python3
"""
PDF Generation using AppleScript and Microsoft Word
"""

import os
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from PyPDF2 import PdfMerger
import tempfile
import shutil
import glob
import random
import subprocess
import time

def replace_variables_in_doc(doc_path, replacements, output_path):
    """Replace variables in the Word document"""
    try:
        doc = Document(doc_path)
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))
        
        # Replace in all tables  
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error: {e}")
        return False

def convert_docx_to_pdf_applescript(docx_path, pdf_path):
    """Convert DOCX to PDF using AppleScript and Microsoft Word"""
    try:
        # AppleScript to convert DOCX to PDF
        applescript = f'''
        on run
            set inputFile to POSIX file "{docx_path}"
            set outputFile to POSIX file "{pdf_path}"
            
            tell application "Microsoft Word"
                activate
                open inputFile
                set activeDoc to active document
                save as activeDoc file name (outputFile as string) file format format PDF
                close activeDoc saving no
            end tell
        end run
        '''
        
        # Run AppleScript
        result = subprocess.run(
            ['osascript', '-e', applescript],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0:
            # Wait a moment for file to be written
            time.sleep(1)
            return os.path.exists(pdf_path)
        else:
            print(f"AppleScript error: {result.stderr}")
            return False
            
    except subprocess.TimeoutExpired:
        print("Conversion timed out")
        return False
    except Exception as e:
        print(f"Conversion error: {e}")
        return False

def process_single_record(row, documents_dir, stamp_papers, output_folder, temp_dir, index, total):
    """Process a single shareholder record"""
    
    # Get details
    name = row.get('Name of the Shareholder', 'Unknown')
    entity_type = row.get('Entity', 'Individual')
    email = row.get('Email', '')
    address = row.get('Address', '')
    
    if pd.isna(entity_type) or entity_type == '':
        entity_type = 'Individual'
    
    print(f"\n[{index}/{total}] {name} ({entity_type})")
    
    # Select template
    if entity_type.lower() == 'entity':
        template_path = os.path.join(documents_dir, 'Entity.docx')
    else:
        template_path = os.path.join(documents_dir, 'Individual.docx')
    
    # Create safe filename
    safe_name = "".join(c for c in str(name) if c.isalnum() or c in (' ', '-', '_')).rstrip()[:50]
    
    # Prepare replacements
    replacements = {
        '[Name of the Shareholder]': str(name),
        '[Email]': str(email),
        '[Address]': str(address),
        '[Date]': str(datetime.now().day)
    }
    
    # Create temp files
    temp_docx = os.path.join(temp_dir, f"{safe_name}.docx")
    temp_pdf = os.path.join(temp_dir, f"{safe_name}.pdf")
    
    # Replace variables
    print("  Replacing variables...")
    if not replace_variables_in_doc(template_path, replacements, temp_docx):
        print("  ✗ Failed")
        return False
    
    # Convert to PDF
    print("  Converting to PDF...")
    if not convert_docx_to_pdf_applescript(temp_docx, temp_pdf):
        print("  ✗ Conversion failed")
        return False
    
    # Merge with stamp paper
    final_pdf = os.path.join(output_folder, f"{safe_name}_{entity_type}.pdf")
    
    try:
        merger = PdfMerger()
        
        # Add stamp paper
        if stamp_papers:
            stamp = random.choice(stamp_papers)
            merger.append(stamp)
        
        # Add document
        merger.append(temp_pdf)
        
        # Save
        merger.write(final_pdf)
        merger.close()
        
        print(f"  ✓ Saved: {os.path.basename(final_pdf)}")
        return True
        
    except Exception as e:
        print(f"  ✗ Merge failed: {e}")
        return False

def main():
    """Main function"""
    
    print("=" * 60)
    print("PDF GENERATION WITH APPLESCRIPT")
    print("=" * 60)
    
    # Setup
    documents_dir = os.path.expanduser("~/Documents/Projects")
    excel_path = os.path.join(documents_dir, "Email + Address List.xlsx")
    stamp_folder = os.path.join(documents_dir, "DOA Stamp Papers")
    output_folder = os.path.join(documents_dir, "Generated_PDFs")
    
    os.makedirs(output_folder, exist_ok=True)
    
    # Read data
    df = pd.read_excel(excel_path)
    stamp_papers = glob.glob(os.path.join(stamp_folder, "*.pdf"))
    
    print(f"\nFound {len(df)} records")
    print(f"Found {len(stamp_papers)} stamp papers")
    
    # Get range
    start = int(input(f"\nStart from record (1-{len(df)}): ") or "1") - 1
    count = int(input("How many to process: ") or "3")
    end = min(start + count, len(df))
    
    print(f"\nProcessing {count} records...")
    print("-" * 60)
    
    # Create temp dir
    temp_dir = tempfile.mkdtemp()
    success = 0
    
    try:
        for i in range(start, end):
            if process_single_record(
                df.iloc[i], 
                documents_dir, 
                stamp_papers, 
                output_folder, 
                temp_dir,
                i - start + 1,
                count
            ):
                success += 1
    
    finally:
        # Cleanup
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    # Summary
    print("\n" + "=" * 60)
    print(f"Processed: {count}")
    print(f"Success: {success}")
    print(f"Failed: {count - success}")
    print(f"Output: {output_folder}")

if __name__ == "__main__":
    main()